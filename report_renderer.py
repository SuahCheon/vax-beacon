"""
Vax-Beacon v4 | Report Renderer (Code Only)
=============================================
Converts pipeline result dict -> Markdown report.
No LLM calls. Reuses Stage 6 officer summary.
Optional .docx conversion with python-docx.
"""

import os
from datetime import datetime

REPORTS_PATH = "reports"


def render_report(result: dict, output_dir: str = REPORTS_PATH) -> str:
    """
    Render a single case result to Markdown.
    Returns the filepath of the generated report.
    """
    os.makedirs(output_dir, exist_ok=True)
    vaers_id = result["vaers_id"]
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")

    if result.get("early_exit"):
        md = _render_early_exit(result, timestamp)
    else:
        md = _render_full(result, timestamp)

    filename = f"VAERS_{vaers_id}_v4.md"
    filepath = os.path.join(output_dir, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(md)

    return filepath


def render_docx(md_path: str) -> str | None:
    """
    Convert Markdown report to Word document.
    Returns filepath or None if python-docx not available.
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  [Report] python-docx not installed, skipping .docx generation")
        return None

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = DocxDocument()

    for line in lines:
        line = line.rstrip("\n")

        # Skip horizontal rules
        if line.strip() == "---":
            continue

        # Headings
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        # Blockquotes -> italic paragraph
        elif line.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            run.italic = True
        # Table rows -> skip header separators, add as plain text
        elif line.startswith("|"):
            if set(line.replace("|", "").replace("-", "").strip()) == set():
                continue  # separator row like |---|---|
            cells = [c.strip() for c in line.split("|")[1:-1]]
            p = doc.add_paragraph("  |  ".join(cells))
            p.style = doc.styles["No Spacing"]
        # Bold lines
        elif line.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("**", "").strip())
            run.bold = True
        # Italic footer
        elif line.startswith("*") and line.endswith("*"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*").strip())
            run.italic = True
            run.font.size = Pt(9)
        # Empty lines
        elif line.strip() == "":
            doc.add_paragraph("")
        # Normal text
        else:
            doc.add_paragraph(line)

    docx_path = md_path.replace(".md", ".docx")
    doc.save(docx_path)
    return docx_path


# ==========================================================================
# Full Pipeline Report (Brighton L1-L3)
# ==========================================================================

def _render_full(result: dict, timestamp: str) -> str:
    """Render full pipeline report (Brighton L1-L3)."""
    s1 = result["stages"].get("stage1_icsr", {})
    s2 = result["stages"].get("stage2_brighton", {})
    s3 = result["stages"].get("stage3_ddx", {})
    s4 = result["stages"].get("stage4_temporal", {})
    s5 = result["stages"].get("stage5_causality", {})
    s6 = result["stages"].get("stage6_guidance", {})

    lines = []
    lines.append(_header(result, timestamp))
    lines.append(_patient_profile(s1))
    lines.append(_brighton(s2))
    lines.append(_ddx(s3))
    lines.append(_temporal(s4))
    lines.append(_classification(s5))
    lines.append(_guidance(s6))
    lines.append(_officer_summary(s6))
    lines.append(_footer(timestamp))

    return "\n".join(lines)


# ==========================================================================
# Early Exit Report (Brighton L4)
# ==========================================================================

def _render_early_exit(result: dict, timestamp: str) -> str:
    """Render early exit report (Brighton L4)."""
    s1 = result["stages"].get("stage1_icsr", {})
    s2 = result["stages"].get("stage2_brighton", {})
    s6 = result["stages"].get("stage6_guidance", {})

    lines = []
    lines.append(_header(result, timestamp))
    lines.append(_patient_profile(s1))
    lines.append(_brighton(s2, early_exit=True))
    lines.append(_early_exit_guidance(s6))
    lines.append(_officer_summary(s6))
    lines.append(_footer(timestamp))

    return "\n".join(lines)


# ==========================================================================
# Section Renderers
# ==========================================================================

def _header(result: dict, timestamp: str) -> str:
    total_time = result.get("processing_time", {}).get("total", "?")
    vaers_id = result["vaers_id"]
    return (
        f"# Vax-Beacon v4.3 — AEFI Causality Assessment Report\n\n"
        f"**VAERS ID:** {vaers_id} | "
        f"**Generated:** {timestamp} | "
        f"**Processing:** {total_time}s\n\n"
        f"---\n"
    )


def _patient_profile(s1: dict) -> str:
    demo = s1.get("demographics", {})
    vax = s1.get("vaccine", {})
    event = s1.get("event", {})
    hist = s1.get("medical_history", {})
    outcomes = s1.get("outcomes", {})
    dq = s1.get("data_quality", {})

    age = demo.get("age", "?")
    sex = demo.get("sex", "?")
    state = demo.get("state", "?")

    vax_name = vax.get("name", "?")
    dose = vax.get("dose_number", "?")
    vax_date = vax.get("vaccination_date", "?")

    onset_date = event.get("onset_date", "?")
    days = event.get("days_to_onset", "?")
    symptoms = ", ".join(event.get("symptoms", [])) or "None reported"

    prior = ", ".join(hist.get("prior_conditions", [])) or "None"
    meds = ", ".join(hist.get("medications", [])) or "None"

    # Outcome string
    outcome_parts = []
    if outcomes.get("hospitalized"):
        outcome_parts.append("Hospitalized")
    if outcomes.get("er_visit"):
        outcome_parts.append("ER visit")
    if outcomes.get("life_threatening"):
        outcome_parts.append("Life-threatening")
    if outcomes.get("died"):
        outcome_parts.append("Died")
    recovered = outcomes.get("recovered", "?")
    if recovered == "Y":
        outcome_parts.append("Recovered")
    elif recovered == "N":
        outcome_parts.append("Not recovered")
    outcome_str = ", ".join(outcome_parts) or "Not specified"

    richness = dq.get("richness_score", "?")

    return (
        f"\n## Patient Profile\n\n"
        f"| Field | Value |\n"
        f"|-------|-------|\n"
        f"| Age / Sex | {age} {sex} |\n"
        f"| State | {state} |\n"
        f"| Vaccine | {vax_name}, Dose #{dose}, {vax_date} |\n"
        f"| Onset | {onset_date} ({days} days post-vaccination) |\n"
        f"| Symptoms | {symptoms} |\n"
        f"| Medical History | {prior} |\n"
        f"| Medications | {meds} |\n"
        f"| Outcome | {outcome_str} |\n"
        f"| Data Quality | Richness {richness}/10 |\n"
    )


def _brighton(s2: dict, early_exit: bool = False) -> str:
    level = s2.get("brighton_level", "?")
    justification = s2.get("brighton_justification", "")
    criteria = s2.get("criteria_met", {})
    pending = s2.get("pending_overrides")

    def _check(key):
        return "✓" if criteria.get(key) else "✗"

    section = (
        f"\n## Diagnostic Validation (Brighton)\n\n"
        f"**Brighton Level: {level}** — {justification}\n\n"
        f"| Criterion | Result |\n"
        f"|-----------|--------|\n"
        f"| Histopathology | {_check('histopathology')} |\n"
        f"| Cardiac MRI | {_check('cardiac_mri_positive')} |\n"
        f"| Troponin elevated | {_check('troponin_elevated')} |\n"
        f"| ECG abnormal | {_check('ecg_abnormal')} |\n"
        f"| Echo abnormal | {_check('echo_abnormal')} |\n"
        f"| Compatible symptoms | {_check('compatible_symptoms')} |\n"
        f"| Inflammatory markers | {_check('inflammatory_markers_elevated')} |\n"
    )

    if pending:
        pending_items = ", ".join(f"{k}: \"{v}\"" for k, v in pending.items())
        section += f"\n**Pending Overrides (v4.3):** {pending_items}\n"

    if early_exit:
        section += (
            f"\n> **EARLY EXIT** — Insufficient diagnostic evidence (Brighton Level 4). "
            f"Stages 3-5 skipped. Proceeding directly to Stage 6 guidance.\n"
        )

    return section


def _ddx(s3: dict) -> str:
    """Render Stage 3 DDx results."""
    lines = ["\n## Differential Diagnosis (WHO Step 1)\n"]

    # 3A observations summary
    obs_data = s3.get("v4_stage3a", {}).get("clinical_observations", {})
    if obs_data:
        lines.append("### Clinical Observations\n")
        for domain, observations in obs_data.items():
            if observations:
                domain_label = domain.replace("_", " ").title()
                lines.append(f"**{domain_label}:**")
                for obs in observations:
                    finding = obs.get("finding", "?")
                    context = obs.get("context", "")
                    conf = obs.get("confidence", "")
                    lines.append(f"- {finding} — *\"{context}\"* [{conf}]")
                lines.append("")

    # Key negatives
    negatives = s3.get("v4_stage3a", {}).get("key_negatives", [])
    if negatives:
        lines.append("**Key Negatives:** " + "; ".join(negatives))
        lines.append("")

    # NCI table
    nci_detailed = s3.get("nci_detailed", {})
    if nci_detailed:
        lines.append("### DDx Candidates\n")
        lines.append("| Etiology | NCI Score | Key Markers | Status |")
        lines.append("|----------|-----------|-------------|--------|")
        for etiology, detail in nci_detailed.items():
            if etiology == "narrative_nuance":
                continue
            nci = detail.get("nci_score", 0)
            if nci == 0:
                continue
            markers = [m["marker"] for m in detail.get("markers_passed", [])]
            markers_str = ", ".join(markers) if markers else "—"
            # Determine status
            if nci >= 0.7:
                status = "CONFIRMED"
            elif nci >= 0.4:
                status = "SUSPECTED"
            else:
                status = "WEAK"
            label = etiology.replace("_", " ").title()
            lines.append(f"| {label} | {nci:.2f} | {markers_str} | {status} |")
        lines.append("")

    # Step 1 conclusion
    step1 = s3.get("who_step1_conclusion", "?")
    dominant = s3.get("dominant_alternative", "NONE")
    max_nci = s3.get("max_nci_score", "?")
    lines.append(f"**Max NCI:** {max_nci}")
    lines.append(f"**WHO Step 1 Conclusion:** {step1}")
    lines.append(f"**Dominant Alternative:** {dominant}")
    lines.append("")

    # Plausibility highlights — show markers that were present with rationale
    markers = s3.get("llm_markers_extracted", {})
    present_markers = {k: v for k, v in markers.items()
                       if v.get("present") and k not in
                       ("reporter_uncertainty", "alternative_suspected", "lack_of_testing")}
    if present_markers:
        lines.append("### Plausibility Highlights\n")
        for marker, detail in present_markers.items():
            plaus = detail.get("plausibility", "?")
            rationale = detail.get("biological_rationale", "")
            lines.append(f"- **{marker}** (plausibility: {plaus}): {rationale}")
        lines.append("")

    # Epistemic uncertainty (narrative nuance)
    nuance = s3.get("narrative_nuance", {})
    nuance_markers = nuance.get("markers_passed", [])
    if nuance_markers:
        lines.append("### Narrative Nuance\n")
        for m in nuance_markers:
            lines.append(f"- **{m['marker']}** (w={m['weight']}): {m.get('rationale', '')}")
        lines.append("")

    # Information gaps
    gaps = s3.get("information_gaps", [])
    if gaps:
        lines.append("### Information Gaps\n")
        for gap in gaps:
            lines.append(f"- {gap}")
        lines.append("")

    return "\n".join(lines)


def _temporal(s4: dict) -> str:
    """Render Stage 4 temporal assessment."""
    known = s4.get("known_ae_assessment", {})
    temporal = s4.get("temporal_assessment", {})

    is_known = known.get("is_known_ae", "?")
    evidence = known.get("evidence_level", "?")
    source = known.get("source", "?")

    days = temporal.get("days_to_onset", "?")
    zone = temporal.get("temporal_zone", "?")
    nam = temporal.get("nam_2024_alignment", "?")
    step2 = s4.get("who_step2_met", "?")

    flags = s4.get("flags", [])
    flags_str = ", ".join(flags) if flags else "None"

    return (
        f"\n## Temporal Assessment (WHO Step 2)\n\n"
        f"| Parameter | Value |\n"
        f"|-----------|-------|\n"
        f"| Known AE | {is_known} ({evidence} — {source}) |\n"
        f"| Days to onset | {days} |\n"
        f"| Temporal zone | {zone} |\n"
        f"| NAM 2024 alignment | {nam} |\n"
        f"| WHO Step 2 Met | {step2} |\n"
        f"| Flags | {flags_str} |\n"
    )


def _classification(s5: dict) -> str:
    """Render Stage 5 classification."""
    who = s5.get("who_category", "?")
    label = s5.get("who_category_label", "")
    confidence = s5.get("confidence", "?")
    source = s5.get("classification_source", "?")
    reasoning = s5.get("reasoning", "")

    chain = s5.get("decision_chain", {})

    # Format decision chain
    chain_lines = []
    if chain:
        chain_lines.append(f"- Q1 Valid diagnosis: {chain.get('q1_valid_diagnosis', '?')} (Brighton L{chain.get('brighton_level', '?')})")
        chain_lines.append(f"- Q2 Definite other cause: {chain.get('q2_definite_other_cause', '?')} (max_nci={chain.get('max_nci', '?')})")
        if chain.get("q3_known_ae") is not None:
            chain_lines.append(f"- Q3 Known AE: {chain.get('q3_known_ae', '?')}")
        if chain.get("q4_temporal_met") is not None:
            chain_lines.append(f"- Q4 Temporal met: {chain.get('q4_temporal_met', '?')} (zone={chain.get('temporal_zone', '?')})")
        if chain.get("q5_conflicting_alternatives") is not None:
            chain_lines.append(f"- Q5 Conflicting alternatives: {chain.get('q5_conflicting_alternatives', '?')}")
    chain_str = "\n".join(chain_lines)

    # Key factors
    factors = s5.get("key_factors", [])
    factors_str = "\n".join(f"- {f}" for f in factors) if factors else "None"

    return (
        f"\n## WHO Causality Classification\n\n"
        f"### **{who}** — {label}\n\n"
        f"**Confidence:** {confidence}\n"
        f"**Classification Source:** {source}\n\n"
        f"**Decision Chain:**\n{chain_str}\n\n"
        f"**Key Factors:**\n{factors_str}\n\n"
        f"**Reasoning:**\n{reasoning}\n"
    )


def _guidance(s6: dict) -> str:
    """Render Stage 6 guidance (normal flow)."""
    risk = s6.get("overall_risk_signal", "?")
    escalation = s6.get("escalation_needed", False)
    escalation_reason = s6.get("escalation_reason", "")

    lines = [f"\n## Investigation Guidance\n"]
    lines.append(f"**Overall Risk Signal:** {risk}")
    if escalation:
        lines.append(f"**Escalation Needed:** Yes — {escalation_reason}")
    lines.append("")

    # Investigative gaps table
    gaps = s6.get("investigative_gaps", [])
    if gaps:
        lines.append("| Priority | Gap | Action | Classification Impact |")
        lines.append("|----------|-----|--------|----------------------|")
        for gap in gaps:
            priority = gap.get("priority", "?")
            gap_text = gap.get("gap", "?")
            action = gap.get("action", "?")
            impact = gap.get("impact_on_classification", "?")
            lines.append(f"| {priority} | {gap_text} | {action} | {impact} |")
        lines.append("")

    # Recommended actions
    actions = s6.get("recommended_actions", [])
    if actions:
        lines.append("### Recommended Actions\n")
        for a in actions:
            lines.append(f"- {a}")
        lines.append("")

    # Quality flags
    qf = s6.get("quality_flags", {})
    if qf:
        lines.append("### Quality Assessment\n")
        lines.append("| Flag | Value |")
        lines.append("|------|-------|")
        for key, val in qf.items():
            label = key.replace("_", " ").title()
            lines.append(f"| {label} | {val} |")
        lines.append("")

    return "\n".join(lines)


def _early_exit_guidance(s6: dict) -> str:
    """Render Stage 6 guidance for Brighton L4 early exit."""
    risk = s6.get("overall_risk_signal", "?")
    who = s6.get("who_category", "Unclassifiable")
    reason = s6.get("unclassifiable_reason", "?")

    lines = [f"\n## Brighton Exit Guidance\n"]
    lines.append(f"**WHO Category:** {who}")
    lines.append(f"**Unclassifiable Reason:** {reason}")
    lines.append(f"**Overall Risk Signal:** {risk}")
    lines.append("")

    # What is known / missing
    known = s6.get("what_is_known", "")
    missing = s6.get("what_is_missing", [])
    if known:
        lines.append("### What is Known\n")
        lines.append(known)
        lines.append("")
    if missing:
        lines.append("### What is Missing\n")
        for m in missing:
            lines.append(f"- {m}")
        lines.append("")

    # Diagnostic deficiencies
    deficiencies = s6.get("diagnostic_deficiencies", [])
    if deficiencies:
        lines.append("### Diagnostic Deficiencies\n")
        lines.append("| Priority | Missing Test | Importance | Action |")
        lines.append("|----------|-------------|------------|--------|")
        for d in deficiencies:
            priority = d.get("priority", "?")
            test = d.get("missing_test", "?")
            importance = d.get("importance", "?")
            action = d.get("action", "?")
            lines.append(f"| {priority} | {test} | {importance} | {action} |")
        lines.append("")

    # Fastest path
    fastest = s6.get("fastest_path_to_classification", {})
    if fastest:
        target = fastest.get("target_level", "?")
        tests = ", ".join(fastest.get("required_tests", []))
        explanation = fastest.get("explanation", "")
        lines.append("### Fastest Path to Classification\n")
        lines.append(f"**Target:** {target}")
        lines.append(f"**Required Tests:** {tests}")
        lines.append(f"**Explanation:** {explanation}")
        lines.append("")

    # Alternative diagnoses
    alt = s6.get("alternative_diagnoses", "")
    if alt:
        lines.append("### Alternative Diagnoses Noted\n")
        lines.append(alt)
        lines.append("")

    # Quality flags
    qf = s6.get("quality_flags", {})
    if qf:
        lines.append("### Quality Assessment\n")
        lines.append("| Flag | Value |")
        lines.append("|------|-------|")
        for key, val in qf.items():
            label = key.replace("_", " ").title()
            lines.append(f"| {label} | {val} |")
        lines.append("")

    # Reassessment potential
    reassess = s6.get("reassessment_potential", "")
    if reassess:
        lines.append("### Reassessment Potential\n")
        lines.append(reassess)
        lines.append("")

    return "\n".join(lines)


def _officer_summary(s6: dict) -> str:
    """Render officer summary as blockquote."""
    summary = s6.get("officer_summary", "No summary available.")
    return (
        f"\n## Officer Summary\n\n"
        f"> {summary}\n"
    )


def _footer(timestamp: str) -> str:
    return (
        f"\n---\n"
        f"*Report generated by Vax-Beacon v4.3 | {timestamp}*\n"
    )
